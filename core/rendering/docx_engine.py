"""GOST-compliant Russian Legal DOCX Document Generation Engine.

Applies strict ГОСТ Р 7.0.97-2016 formatting, XML font bindings (w:rFonts),
anti-orphan controls (w:widowControl, w:cantSplit, w:tblHeader), and dual-column signature blocks.
"""

from __future__ import annotations
import os
import io
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from core.models.base import BaseContract
from core.models.supply import SupplyContract
from core.models.services import ServiceContract
from core.models.work import WorkContract
from core.templates.registry import ContractRegistry
from core.num_to_words import format_rubles, amount_to_words_ru


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell internal margins (padding) in dxa (1 mm = ~56.7 dxa)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_row_cant_split(row):
    """Prevent table row from breaking across page boundaries."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def _set_row_header(row):
    """Mark row as repeating table header on page split."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def _apply_rfonts(run, font_name="Liberation Serif"):
    """Explicitly set w:rFonts attribute for all Unicode scripts (ascii, hAnsi, cs, eastAsia)."""
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)


def _set_paragraph_gost_format(p, indent=True, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Apply standard Russian GOST formatting to paragraph."""
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Mm(12.5)  # 1.25 cm красная строка

    # Set widow/orphan control on XML level
    pPr = p._p.get_or_add_pPr()
    widow = OxmlElement("w:widowControl")
    pPr.append(widow)


class DocxEngine:
    """Deterministic DOCX document generator with GOST formatting."""

    FONT_NAME = "Liberation Serif"

    @classmethod
    def generate(cls, contract: BaseContract) -> io.BytesIO:
        """Generate formatted DOCX document and return as BytesIO stream."""
        doc_structure = ContractRegistry.assemble_document_structure(contract)

        doc = Document()

        # 1. Page Margins (ГОСТ Р 7.0.97-2016: Левое 25 мм, Правое 15 мм, Верх 20 мм, Низ 20 мм)
        section = doc.sections[0]
        section.left_margin = Mm(25)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        # Base style config
        normal_style = doc.styles["Normal"]
        normal_style.font.name = cls.FONT_NAME
        normal_style.font.size = Pt(11.5)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)

        header = doc_structure["header"]
        meta = doc_structure["metadata"]

        # Document Title
        p_title = doc.add_paragraph()
        _set_paragraph_gost_format(p_title, indent=False, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        r_title = p_title.add_run(header["title"])
        r_title.bold = True
        r_title.font.size = Pt(13)
        _apply_rfonts(r_title, cls.FONT_NAME)

        # City and Date header line (table or two-sided text)
        p_meta = doc.add_paragraph()
        _set_paragraph_gost_format(p_meta, indent=False, space_after=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r_city = p_meta.add_run(header["city"])
        _apply_rfonts(r_city, cls.FONT_NAME)

        # Tab to right side for date
        p_meta.paragraph_format.tab_stops.add_tab_stop(Mm(170))
        r_tab = p_meta.add_run("\t")
        r_date = p_meta.add_run(header["date"])
        _apply_rfonts(r_date, cls.FONT_NAME)

        # Preamble text
        p_pre = doc.add_paragraph()
        _set_paragraph_gost_format(p_pre, indent=True, space_after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r_pre = p_pre.add_run(header["text"])
        _apply_rfonts(r_pre, cls.FONT_NAME)

        # Sections & Clauses
        for section_data in doc_structure["sections"]:
            # Section Heading
            p_sec = doc.add_paragraph()
            _set_paragraph_gost_format(p_sec, indent=False, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
            p_sec.paragraph_format.keep_with_next = True
            r_sec = p_sec.add_run(f"{section_data['section_num']}. {section_data['title']}")
            r_sec.bold = True
            r_sec.font.size = Pt(12)
            _apply_rfonts(r_sec, cls.FONT_NAME)

            # Paragraphs
            for clause in section_data["clauses"]:
                p_cl = doc.add_paragraph()
                _set_paragraph_gost_format(p_cl, indent=True, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                r_num = p_cl.add_run(f"{clause['num']}. ")
                r_num.bold = True
                _apply_rfonts(r_num, cls.FONT_NAME)

                r_txt = p_cl.add_run(clause["text"])
                _apply_rfonts(r_txt, cls.FONT_NAME)

            # If this is Section 1 or 2 and we have tables (Supply Items, Work Stages, Services), render clean table
            if section_data["section_num"] == "1":
                cls._render_contract_specific_tables(doc, contract)

        # Signatures Block
        cls._render_signatures_block(doc, doc_structure["signatures"])

        # Output to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @classmethod
    def _render_contract_specific_tables(cls, doc: Document, contract: BaseContract):
        """Render specification tables for Supply, Work, and Services."""
        if isinstance(contract, SupplyContract) and contract.items:
            # Specification Table
            table = doc.add_table(rows=1, cols=6)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            headers = ["№", "Наименование товара", "Ед. изм.", "Кол-во", "Цена (руб.)", "Сумма (руб.)"]
            widths = [Mm(10), Mm(75), Mm(18), Mm(18), Mm(24), Mm(25)]

            # Header Row
            hdr_row = table.rows[0]
            _set_row_header(hdr_row)
            _set_row_cant_split(hdr_row)

            for i, title in enumerate(headers):
                cell = hdr_row.cells[i]
                cell.width = widths[i]
                _set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
                p = cell.paragraphs[0]
                _set_paragraph_gost_format(p, indent=False, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(10)
                _apply_rfonts(r, cls.FONT_NAME)
                # Background fill for header
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFEFEF"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Data Rows
            for idx, item in enumerate(contract.items, start=1):
                row = table.add_row()
                _set_row_cant_split(row)

                data = [
                    (str(idx), WD_ALIGN_PARAGRAPH.CENTER),
                    (item.name, WD_ALIGN_PARAGRAPH.LEFT),
                    (item.unit, WD_ALIGN_PARAGRAPH.CENTER),
                    (f"{item.quantity:g}", WD_ALIGN_PARAGRAPH.RIGHT),
                    (format_rubles(item.price_per_unit), WD_ALIGN_PARAGRAPH.RIGHT),
                    (format_rubles(item.total_price), WD_ALIGN_PARAGRAPH.RIGHT),
                ]

                for col_idx, (text, align) in enumerate(data):
                    cell = row.cells[col_idx]
                    cell.width = widths[col_idx]
                    _set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    p = cell.paragraphs[0]
                    _set_paragraph_gost_format(p, indent=False, space_after=0, align=align)
                    r = p.add_run(text)
                    r.font.size = Pt(9.5)
                    _apply_rfonts(r, cls.FONT_NAME)

            # Total Row
            tot_row = table.add_row()
            _set_row_cant_split(tot_row)
            cell_lbl = tot_row.cells[0]
            # Merge 0 to 4
            for c in range(1, 5):
                cell_lbl.merge(tot_row.cells[c])
            _set_cell_margins(cell_lbl, top=100, bottom=100, left=100, right=100)
            p = cell_lbl.paragraphs[0]
            _set_paragraph_gost_format(p, indent=False, space_after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
            r = p.add_run("ИТОГО:")
            r.bold = True
            r.font.size = Pt(10)
            _apply_rfonts(r, cls.FONT_NAME)

            cell_val = tot_row.cells[5]
            _set_cell_margins(cell_val, top=100, bottom=100, left=100, right=100)
            p_val = cell_val.paragraphs[0]
            _set_paragraph_gost_format(p_val, indent=False, space_after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
            r_val = p_val.add_run(format_rubles(contract.total_amount))
            r_val.bold = True
            r_val.font.size = Pt(10)
            _apply_rfonts(r_val, cls.FONT_NAME)

            # Apply borders to table
            cls._apply_table_borders(table)

            # Spacing after table
            p_after = doc.add_paragraph()
            _set_paragraph_gost_format(p_after, indent=False, space_after=6)

        elif isinstance(contract, WorkContract) and contract.stages:
            # Stages Table
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            headers = ["Этап", "Наименование работ", "Сроки выполнения", "Результат", "Стоимость (руб.)"]
            widths = [Mm(15), Mm(60), Mm(35), Mm(35), Mm(25)]

            hdr_row = table.rows[0]
            _set_row_header(hdr_row)
            _set_row_cant_split(hdr_row)

            for i, title in enumerate(headers):
                cell = hdr_row.cells[i]
                cell.width = widths[i]
                _set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
                p = cell.paragraphs[0]
                _set_paragraph_gost_format(p, indent=False, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
                r = p.add_run(title)
                r.bold = True
                r.font.size = Pt(10)
                _apply_rfonts(r, cls.FONT_NAME)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFEFEF"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            for stage in contract.stages:
                row = table.add_row()
                _set_row_cant_split(row)
                data = [
                    (f"№ {stage.stage_number}", WD_ALIGN_PARAGRAPH.CENTER),
                    (stage.title, WD_ALIGN_PARAGRAPH.LEFT),
                    (f"с {stage.start_date}\nпо {stage.end_date}", WD_ALIGN_PARAGRAPH.CENTER),
                    (stage.deliverable_result or "—", WD_ALIGN_PARAGRAPH.LEFT),
                    (format_rubles(stage.cost), WD_ALIGN_PARAGRAPH.RIGHT),
                ]
                for col_idx, (text, align) in enumerate(data):
                    cell = row.cells[col_idx]
                    cell.width = widths[col_idx]
                    _set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    p = cell.paragraphs[0]
                    _set_paragraph_gost_format(p, indent=False, space_after=0, align=align)
                    r = p.add_run(text)
                    r.font.size = Pt(9.5)
                    _apply_rfonts(r, cls.FONT_NAME)

            cls._apply_table_borders(table)
            p_after = doc.add_paragraph()
            _set_paragraph_gost_format(p_after, indent=False, space_after=6)

    @classmethod
    def _apply_table_borders(cls, table):
        """Add clean subtle borders to table."""
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="DCDCDC"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    @classmethod
    def _render_signatures_block(cls, doc: Document, signatures_data: Dict[str, Any]):
        """Render dual-column Russian requisites and signature block."""
        # Heading
        p_head = doc.add_paragraph()
        _set_paragraph_gost_format(p_head, indent=False, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        p_head.paragraph_format.keep_with_next = True
        r_head = p_head.add_run("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
        r_head.bold = True
        r_head.font.size = Pt(12)
        _apply_rfonts(r_head, cls.FONT_NAME)

        # 2-column table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        row = table.rows[0]
        _set_row_cant_split(row)

        col_widths = [Mm(85), Mm(85)]

        client_info = signatures_data["client"]
        vendor_info = signatures_data["vendor"]

        parties = [client_info, vendor_info]

        for col_idx, party_info in enumerate(parties):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            _set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

            # Role Title
            p_role = cell.paragraphs[0]
            _set_paragraph_gost_format(p_role, indent=False, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
            r_role = p_role.add_run(party_info["role_title"])
            r_role.bold = True
            r_role.font.size = Pt(10.5)
            _apply_rfonts(r_role, cls.FONT_NAME)

            # Requisites lines
            for line in party_info["details_lines"]:
                p_line = cell.add_paragraph()
                _set_paragraph_gost_format(p_line, indent=False, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
                r_l = p_line.add_run(line)
                r_l.font.size = Pt(9.5)
                _apply_rfonts(r_l, cls.FONT_NAME)

            # Signatory line
            p_sig = cell.add_paragraph()
            _set_paragraph_gost_format(p_sig, indent=False, space_after=20, align=WD_ALIGN_PARAGRAPH.LEFT)
            p_sig.paragraph_format.space_before = Pt(12)
            r_sig_pos = p_sig.add_run(f"{party_info['signatory_position']}:\n\n")
            r_sig_pos.font.size = Pt(9.5)
            _apply_rfonts(r_sig_pos, cls.FONT_NAME)

            r_sig_line = p_sig.add_run("__________________ / " + party_info["signatory_name"] + " /\n")
            r_sig_line.font.size = Pt(9.5)
            _apply_rfonts(r_sig_line, cls.FONT_NAME)

            r_stamp = p_sig.add_run("М.П.")
            r_stamp.font.size = Pt(9.5)
            r_stamp.bold = True
            _apply_rfonts(r_stamp, cls.FONT_NAME)

        # Borderless signature table
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none"/>'
            f'<w:bottom w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
